#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
采购文件填充工具 - 命令行版本
功能：将Excel信息表中的数据填充到Word模板中

用法：
    python3 procurement_filler_cli.py <信息表.xlsx> <模板.docx> <输出.docx>
"""

import sys
import os
import re
import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import copy


def read_excel_data(excel_path):
    """读取 Excel 信息表，返回数据字典"""
    print(f"📖 读取 Excel: {excel_path}")
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    
    # 读取信息表
    if "信息表" not in wb.sheetnames:
        raise ValueError(f"Excel 中找不到 '信息表' 工作表！可用的工作表: {wb.sheetnames}")
    
    ws = wb["信息表"]
    data_dict = {}
    
    # 查找列索引
    headers = {}
    for col in range(1, ws.max_column + 1):
        header = ws.cell(1, col).value
        if header:
            headers[header] = col
    
    if "数据名" not in headers or "数据信息" not in headers:
        raise ValueError("Excel 信息表缺少 '数据名' 或 '数据信息' 列！")
    
    name_col = headers["数据名"]
    value_col = headers["数据信息"]
    
    # 读取数据
    for row in range(2, ws.max_row + 1):
        name = ws.cell(row, name_col).value
        value = ws.cell(row, value_col).value
        if name:
            # 修复问题1: 处理科学计数法 - 将数字转为字符串，避免科学计数法
            if isinstance(value, (int, float)):
                # 如果是整数或看起来像银行账号的数字，转为字符串
                if isinstance(value, int) or (isinstance(value, float) and value == int(value)):
                    value = str(int(value))
                else:
                    value = str(value)
            elif value is None:
                value = ""
            else:
                value = str(value)
            
            data_dict[name] = value
    
    # 添加别名映射：【叁】-> 采购年限的值
    if "采购年限" in data_dict and "叁" not in data_dict:
        data_dict["叁"] = data_dict["采购年限"]
        print(f"  添加别名: 【叁】 = {data_dict['叁']}")
    
    print(f"✓ 共读取 {len(data_dict)} 个数据项")
    return data_dict


def convert_doc_to_docx(doc_path):
    """使用 LibreOffice 将 .doc 转换为 .docx"""
    print(f"🔄 转换 .doc 到 .docx: {doc_path}")
    import tempfile
    
    temp_dir = tempfile.gettempdir()
    filename = os.path.basename(doc_path)
    
    # 执行转换
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'docx', doc_path, '--outdir', temp_dir],
        capture_output=True, text=True
    )
    
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice 转换失败: {result.stderr}")
    
    # 返回转换后的文件路径
    converted_path = os.path.join(temp_dir, filename.replace('.doc', '.docx').replace('.DOC', '.docx'))
    print(f"✓ 转换完成: {converted_path}")
    return converted_path


def replace_text_in_paragraph(para, data_dict):
    """在段落中替换占位符，处理跨run的情况"""
    replaced_count = 0
    not_found = []
    
    # 获取段落的完整文本
    full_text = para.text
    
    # 查找所有【xxx】格式的占位符
    placeholders = re.findall(r'【([^】]+)】', full_text)
    
    for placeholder in placeholders:
        if placeholder in data_dict:
            old_text = f"【{placeholder}】"
            new_text = data_dict[placeholder]
            full_text = full_text.replace(old_text, new_text)
            replaced_count += 1
        else:
            if placeholder not in not_found:
                not_found.append(placeholder)
    
    # 如果文本发生了变化，需要重新设置run的文本
    if placeholders:
        # 清空所有runs的文本
        for run in para.runs:
            run.text = ""
        # 将新文本放入第一个run
        if para.runs:
            para.runs[0].text = full_text
    
    return replaced_count, not_found


def replace_placeholders(doc, data_dict):
    """替换文档中的占位符【xxx】"""
    print("\n📝 替换文本占位符...")
    total_replaced = 0
    all_not_found = []
    
    # 遍历所有段落
    for para in doc.paragraphs:
        count, not_found = replace_text_in_paragraph(para, data_dict)
        total_replaced += count
        all_not_found.extend(not_found)
    
    # 遍历所有表格中的单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count, not_found = replace_text_in_paragraph(para, data_dict)
                    total_replaced += count
                    all_not_found.extend(not_found)
    
    # 去重未找到的占位符
    all_not_found = list(set(all_not_found))
    
    print(f"✓ 替换了 {total_replaced} 个占位符")
    if all_not_found:
        print(f"⚠ 未找到匹配的占位符 ({len(all_not_found)} 个): {', '.join(all_not_found[:10])}{'...' if len(all_not_found) > 10 else ''}")
    
    return total_replaced


def copy_cell_border(source_cell, target_cell):
    """复制单元格的边框样式"""
    # 获取源单元格的属性
    source_tc = source_cell._tc
    target_tc = target_cell._tc
    
    # 获取或创建 tcPr (table cell properties)
    source_tcPr = source_tc.get_or_add_tcPr()
    target_tcPr = target_tc.get_or_add_tcPr()
    
    # 查找源单元格的边框定义
    source_tcBorders = source_tcPr.first_child_found_in('w:tcBorders')
    if source_tcBorders is not None:
        # 复制边框到新单元格
        new_tcBorders = copy.deepcopy(source_tcBorders)
        # 移除目标单元格现有的边框定义
        existing_borders = target_tcPr.first_child_found_in('w:tcBorders')
        if existing_borders is not None:
            target_tcPr.remove(existing_borders)
        target_tcPr.append(new_tcBorders)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 移除现有的边框
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    else:
        # 清空现有边框
        for child in list(tcBorders):
            tcBorders.remove(child)
    
    # 设置各边边框
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = OxmlElement(tag)
                element.set(qn('w:val'), 'single')  # 实线
                element.set(qn('w:sz'), '4')  # 边框宽度
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), '000000')  # 黑色
                tcBorders.append(element)


def fill_tables(doc, excel_path):
    """填充采购表格"""
    print("\n📊 填充采购表格...")
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    
    if "采购物品价格" not in wb.sheetnames:
        print("⚠ 未找到 '采购物品价格' 工作表，跳过表格填充")
        return 0
    
    ws = wb["采购物品价格"]
    
    # 读取采购数据（跳过空行）
    items = []
    for row in range(2, ws.max_row + 1):
        values = [ws.cell(row, col).value for col in range(1, 8)]
        if values[1] and str(values[1]).strip():  # 第2列是序号
            items.append({
                '序号': str(values[1]) if values[1] else '',
                '品名': str(values[2]) if values[2] else '',
                '规格': str(values[3]) if values[3] else '',
                '单位': str(values[4]) if values[4] else '',
                '品牌': str(values[5]) if values[5] else '',
                '单价': str(values[6]) if values[6] else ''
            })
    
    print(f"✓ 从 Excel 读取到 {len(items)} 条采购数据")
    
    if not items:
        print("⚠ 没有采购数据需要填充")
        return 0
    
    # 查找采购表格（通常是6列的表格）
    target_table = None
    for table in doc.tables:
        if len(table.columns) == 6:
            header = [cell.text.strip() for cell in table.rows[0].cells]
            if any(h in ['序号', '品名', '规格'] for h in header):
                target_table = table
                break
    
    if not target_table:
        print("⚠ 未找到采购清单表格（6列表格）")
        return 0
    
    print(f"✓ 找到采购表格，当前 {len(target_table.rows)} 行")
    
    # 获取第一行数据行的边框样式作为参考
    reference_row = None
    if len(target_table.rows) > 1:
        reference_row = target_table.rows[1]
    
    # 方案：复用现有行或添加新行
    # 保留表头（第1行），从第2行开始填充
    existing_data_rows = len(target_table.rows) - 1  # 除去表头
    
    # 填充或添加行
    for i, item in enumerate(items):
        if i < existing_data_rows:
            # 复用现有行
            row = target_table.rows[i + 1]
        else:
            # 添加新行
            row = target_table.add_row()
            # 修复问题2: 复制边框样式
            if reference_row:
                for j, cell in enumerate(row.cells):
                    if j < len(reference_row.cells):
                        set_cell_border(cell, top=True, bottom=True, left=True, right=True)
        
        cells = row.cells
        cells[0].text = item['序号']
        cells[1].text = item['品名']
        cells[2].text = item['规格']
        cells[3].text = item['单位']
        cells[4].text = item['品牌']
        cells[5].text = item['单价']
        
        # 居中对齐
        for cell in cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 如果有多余的空行，清空它们
    total_rows_needed = len(items) + 1  # +1 for header
    while len(target_table.rows) > total_rows_needed:
        row_to_delete = target_table.rows[-1]
        row_to_delete._element.getparent().remove(row_to_delete._element)
    
    print(f"✓ 已填充 {len(items)} 行采购数据")
    return len(items)


def add_page_breaks(doc):
    """修复问题4: 在每一大部分前添加分页符"""
    print("\n📄 添加分页符...")
    
    # 查找需要分页的关键词
    section_patterns = [
        r'第[一二三四五六七八九十]+部分',
        r'第一部分',
        r'第二部分', 
        r'第三部分',
        r'第1部分',
        r'第2部分',
        r'第3部分',
    ]
    
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        # 检查是否是部分标题
        for pattern in section_patterns:
            if re.match(pattern, text):
                # 在这个段落前添加分页符
                run = para.runs[0] if para.runs else para.add_run()
                run._element.getparent().insert(
                    run._element.getparent().index(run._element),
                    OxmlElement('w:br', {qn('w:type'): 'page'})
                )
                count += 1
                print(f"  在 '{text[:30]}...' 前添加分页符")
                break
    
    print(f"✓ 添加了 {count} 个分页符")


def main():
    print("=" * 60)
    print("       采购文件填充工具 - 命令行版本")
    print("=" * 60)
    
    # 检查参数
    if len(sys.argv) != 4:
        print("\n用法:")
        print(f"  python3 {sys.argv[0]} <信息表.xlsx> <模板.docx> <输出.docx>")
        print("\n示例:")
        print(f"  python3 {sys.argv[0]} 信息表.xlsx 合同模板.docx 生成的合同.docx")
        sys.exit(1)
    
    excel_path = sys.argv[1]
    template_path = sys.argv[2]
    output_path = sys.argv[3]
    
    # 检查文件是否存在
    if not os.path.exists(excel_path):
        print(f"✗ 错误: Excel 文件不存在: {excel_path}")
        sys.exit(1)
    
    if not os.path.exists(template_path):
        print(f"✗ 错误: 模板文件不存在: {template_path}")
        sys.exit(1)
    
    try:
        # 步骤 1: 读取 Excel
        data_dict = read_excel_data(excel_path)
        
        # 步骤 2: 读取模板（转换 .doc 如果需要）
        if template_path.lower().endswith('.doc'):
            template_path = convert_doc_to_docx(template_path)
        
        doc = Document(template_path)
        
        # 步骤 3: 替换占位符
        replace_placeholders(doc, data_dict)
        
        # 步骤 4: 填充表格
        fill_tables(doc, excel_path)
        
        # 步骤 5: 添加分页符
        add_page_breaks(doc)
        
        # 保存文件
        doc.save(output_path)
        
        print("\n" + "=" * 60)
        print("✓ 处理完成！")
        print(f"✓ 文件已保存至: {output_path}")
        print("=" * 60)
        
    except Exception as e:
        print(f"\n✗ 错误: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
